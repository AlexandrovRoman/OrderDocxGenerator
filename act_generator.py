import re
from datetime import datetime
from os.path import join
import docx
from box import Box
from docxtpl import DocxTemplate
from config import TASKS_FOLDER, TEMPLATES_FOLDER, CONTRACT_NUMBER, CONTRACT_SIGN_AT, USERNAME
from utils import remove_row, month2str, doc2pdf


def get_context_from_order(filename):
    context = Box()
    doc = docx.Document(filename)
    date_format = "«%d» {} %Y г."

    text = "\n".join(par.text for par in doc.paragraphs)
    order_date_in_string = re.search("\d\d\\.\d\d\\.\d\d\d\d", text).group()
    order_date = datetime.strptime(order_date_in_string, "%d.%m.%Y")
    context.order_date = order_date.strftime(date_format.format(month2str(order_date.month)))
    today = datetime.today()
    context.today = today.strftime(date_format.format(month2str(today.month)))
    context.contract_number = CONTRACT_NUMBER
    context.contract_sign_at = CONTRACT_SIGN_AT.strftime(date_format.format(month2str(CONTRACT_SIGN_AT.month)))
    context.username = USERNAME

    tasks = doc.tables[0]
    keys = ("index", "description", "start", "end", "time", "mui", "price", "total")
    rows = iter(tasks.rows)
    next(rows)  # skip table headers
    context.tasks = [dict(zip(keys, (cell.text for cell in row.cells))) for row in rows]

    return context


def create_act(order_number):
    target_folder = join(TASKS_FOLDER, order_number)

    try:
        context = get_context_from_order(join(TASKS_FOLDER, order_number, f"ЗН №{order_number}.docx"))
        context.order_number = order_number
    except FileNotFoundError:
        print("Заказа с данным номером нет")
        return

    doc = DocxTemplate(join(TEMPLATES_FOLDER, "Act.docx"))
    doc.render(context)
    remove_row(doc.get_docx().tables[0], 1)
    filename = join(target_folder, f"Акт №{order_number}.docx")
    doc.save(filename)
    doc2pdf(filename, join(target_folder, f"Акт №{order_number}.pdf"))


if __name__ == '__main__':
    create_act(input("Введите номер ЗН: "))
