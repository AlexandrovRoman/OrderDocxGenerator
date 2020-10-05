import re
from datetime import datetime
import docx
from box import Box
from config import CONTRACT_SIGN_AT, CONTRACT_NUMBER


def month2str(month_index: int,
              months=('январь', 'февраль', 'март', 'апрель', 'май', 'июнь',
                      'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь')):
    if month_index not in range(1, 13):
        raise ValueError("Incorrect month index")
    return months[month_index - 1]


def get_context_from_docx(filename):
    context = Box()
    doc = docx.Document(filename)
    date_format = "«%d» {} %Y г."

    text = "\n".join(par.text for par in doc.paragraphs)
    order_date_in_string = re.search("\d\d\\.\d\d\\.\d\d\d\d", text).group()
    order_date = datetime.strptime(order_date_in_string, "%d.%m.%Y")
    context.order_date = order_date.strftime(date_format.format(month2str(order_date.month)))
    today = datetime.today()
    context.today = today.strftime(date_format.format(month2str(today.month)))
    context.contract_number = CONTRACT_NUMBER
    context.contract_sign_at = CONTRACT_SIGN_AT.strftime(date_format.format(month2str(CONTRACT_SIGN_AT.month)))

    tasks = doc.tables[0]
    keys = ("index", "description", "start", "end", "time", "mui", "price", "total")
    rows = iter(tasks.rows)
    next(rows)  # skip table headers
    context.tasks = [dict(zip(keys, (cell.text for cell in row.cells))) for row in rows]

    return context


def remove_row(table, index):
    row = table.rows[index]
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)
